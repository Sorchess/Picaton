"""
Сервис транскрибации документов в текст.

Поддерживаемые форматы:
- PDF (.pdf) — извлечение текста из всех страниц
- DOCX (.docx) — извлечение текста из параграфов и таблиц
- TXT (.txt) — чтение как есть (UTF-8, Windows-1251, Latin-1)
- RTF (.rtf) — конвертация RTF → plain text

Используется для автозаполнения поля «О себе» из резюме
или других документов пользователя.
"""

from __future__ import annotations

import io
import re
import logging
from dataclasses import dataclass
from enum import Enum

logger = logging.getLogger(__name__)

# ── Константы ──────────────────────────────────────────────────────────

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_OUTPUT_CHARS = 5000  # Максимум символов в ответе (bio ≤ 2000, но даём запас)

ALLOWED_MIME_TYPES: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "text/rtf": "rtf",
    "application/rtf": "rtf",
    # Некоторые браузеры отдают другие MIME для DOCX
    "application/msword": "docx",
}

ALLOWED_EXTENSIONS: set[str] = {"pdf", "docx", "doc", "txt", "rtf"}


# ── Модели ─────────────────────────────────────────────────────────────


class DocumentFormat(str, Enum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    RTF = "rtf"


@dataclass(frozen=True)
class TranscriptionResult:
    """Результат транскрибации документа."""

    text: str
    original_filename: str
    format: DocumentFormat
    char_count: int
    was_truncated: bool


# ── Исключения ─────────────────────────────────────────────────────────


class DocumentTranscriptionError(Exception):
    """Базовая ошибка транскрибации."""


class UnsupportedFormatError(DocumentTranscriptionError):
    """Неподдерживаемый формат файла."""


class FileTooLargeError(DocumentTranscriptionError):
    """Файл слишком большой."""


class EmptyDocumentError(DocumentTranscriptionError):
    """Документ пуст или не содержит текста."""


class ParseError(DocumentTranscriptionError):
    """Ошибка парсинга документа."""


# ── Сервис ─────────────────────────────────────────────────────────────


class DocumentTranscriberService:
    """Извлекает текст из загруженных документов."""

    def __init__(
        self,
        max_file_size: int = MAX_FILE_SIZE_BYTES,
        max_output_chars: int = MAX_OUTPUT_CHARS,
    ) -> None:
        self._max_file_size = max_file_size
        self._max_output_chars = max_output_chars

    # ── Публичный API ──────────────────────────────────────────────

    async def transcribe(
        self,
        file_content: bytes,
        filename: str,
        content_type: str | None = None,
    ) -> TranscriptionResult:
        """
        Извлечь текст из файла.

        Args:
            file_content: Байты файла.
            filename: Оригинальное имя файла (для определения формата).
            content_type: MIME-тип (опционально, для дополнительной валидации).

        Returns:
            TranscriptionResult с извлечённым текстом.

        Raises:
            FileTooLargeError: Файл превышает допустимый размер.
            UnsupportedFormatError: Неподдерживаемый формат.
            EmptyDocumentError: Документ пуст.
            ParseError: Ошибка при разборе документа.
        """
        # 1. Валидация размера
        if len(file_content) > self._max_file_size:
            size_mb = len(file_content) / (1024 * 1024)
            max_mb = self._max_file_size / (1024 * 1024)
            raise FileTooLargeError(
                f"Файл слишком большой ({size_mb:.1f} МБ). "
                f"Максимальный размер: {max_mb:.0f} МБ."
            )

        # 2. Определение формата
        fmt = self._detect_format(filename, content_type)

        # 3. Извлечение текста
        raw_text = await self._extract_text(file_content, fmt)

        # 4. Постобработка
        cleaned = self._clean_text(raw_text)

        if not cleaned.strip():
            raise EmptyDocumentError(
                "Не удалось извлечь текст из документа. "
                "Возможно, файл содержит только изображения или пуст."
            )

        # 5. Обрезка при необходимости
        was_truncated = len(cleaned) > self._max_output_chars
        if was_truncated:
            cleaned = cleaned[: self._max_output_chars].rsplit(" ", 1)[0] + "…"

        return TranscriptionResult(
            text=cleaned,
            original_filename=filename,
            format=fmt,
            char_count=len(cleaned),
            was_truncated=was_truncated,
        )

    # ── Определение формата ────────────────────────────────────────

    def _detect_format(
        self, filename: str, content_type: str | None
    ) -> DocumentFormat:
        """Определить формат по расширению и MIME-типу."""
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        if ext not in ALLOWED_EXTENSIONS:
            # Попробуем по MIME
            if content_type and content_type in ALLOWED_MIME_TYPES:
                ext = ALLOWED_MIME_TYPES[content_type]
            else:
                raise UnsupportedFormatError(
                    f"Формат «.{ext}» не поддерживается. "
                    f"Поддерживаемые форматы: PDF, DOCX, TXT, RTF."
                )

        format_map: dict[str, DocumentFormat] = {
            "pdf": DocumentFormat.PDF,
            "docx": DocumentFormat.DOCX,
            "doc": DocumentFormat.DOCX,  # Попытаемся прочитать как DOCX
            "txt": DocumentFormat.TXT,
            "rtf": DocumentFormat.RTF,
        }

        return format_map.get(ext, DocumentFormat.TXT)

    # ── Извлечение текста по формату ───────────────────────────────

    async def _extract_text(
        self, content: bytes, fmt: DocumentFormat
    ) -> str:
        """Маршрутизатор к нужному парсеру."""
        extractors = {
            DocumentFormat.PDF: self._extract_pdf,
            DocumentFormat.DOCX: self._extract_docx,
            DocumentFormat.TXT: self._extract_txt,
            DocumentFormat.RTF: self._extract_rtf,
        }

        extractor = extractors.get(fmt)
        if not extractor:
            raise UnsupportedFormatError(f"Неизвестный формат: {fmt}")

        try:
            return extractor(content)
        except (UnsupportedFormatError, EmptyDocumentError):
            raise
        except Exception as e:
            logger.exception("Ошибка парсинга документа формата %s", fmt)
            raise ParseError(
                f"Не удалось прочитать документ. "
                f"Проверьте, что файл не повреждён. Ошибка: {e}"
            ) from e

    # ── PDF ─────────────────────────────────────────────────────────

    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        """Извлечь текст из PDF."""
        from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(content))
        pages_text: list[str] = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text.strip())

        return "\n\n".join(pages_text)

    # ── DOCX ───────────────────────────────────────────────────────

    @staticmethod
    def _extract_docx(content: bytes) -> str:
        """Извлечь текст из DOCX (параграфы + таблицы)."""
        from docx import Document

        doc = Document(io.BytesIO(content))
        parts: list[str] = []

        # Параграфы
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Таблицы (резюме часто содержат таблицы)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    # ── TXT ─────────────────────────────────────────────────────────

    @staticmethod
    def _extract_txt(content: bytes) -> str:
        """Прочитать текстовый файл с автоопределением кодировки."""
        for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
            try:
                return content.decode(encoding)
            except (UnicodeDecodeError, ValueError):
                continue

        # Fallback: игнорируем непонятные символы
        return content.decode("utf-8", errors="ignore")

    # ── RTF ─────────────────────────────────────────────────────────

    @staticmethod
    def _extract_rtf(content: bytes) -> str:
        """Конвертировать RTF → plain text."""
        from striprtf.striprtf import rtf_to_text

        raw = content.decode("utf-8", errors="ignore")
        return rtf_to_text(raw)

    # ── Постобработка текста ───────────────────────────────────────

    @staticmethod
    def _clean_text(text: str) -> str:
        """
        Очистить извлечённый текст:
        - Убрать лишние пробелы и переносы
        - Нормализовать Unicode-пробелы
        - Удалить непечатаемые символы
        """
        # Удалить непечатаемые символы (кроме переноса строки и табуляции)
        text = re.sub(r"[^\S\n\t]+", " ", text)
        # Убрать множественные пустые строки (оставить максимум одну)
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Убрать пробелы в начале/конце строк
        text = "\n".join(line.strip() for line in text.splitlines())
        # Убрать пустые строки в начале и конце
        text = text.strip()

        return text
